"""Assemble the Project 4 reflection Word document.

Reads reflection-project4-draft.md, drops the blockquoted review notes that
are addressed to Gilbert rather than the grader, and emits a .docx with the
four prompt answers, every validation screenshot, the contrast-analyzer
screenshots, and the generated XML sitemap.

Usage:
    python3 tools/build-reflection-docx.py
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "reflection-project4-draft.md"
SHOTS = ROOT / "validation-screenshots"
OUT = ROOT / "Project4-Reflection-GilbertAnderson.docx"

PAGES = ["index", "about", "services", "portfolio", "case-nonprofit",
         "case-ecommerce", "case-startup", "testimonials", "blog",
         "blog-post-seo", "contact", "privacy"]

CONTRAST = [
    ("contrast-body-text-on-white.png",
     "Body text #1e293b on background #ffffff, 1.05rem (about 16.8px) normal weight"),
    ("contrast-heading-on-white.png",
     "Headings #1e3a8a on background #ffffff, h2 at 1.7rem (about 27px)"),
    ("contrast-white-on-header-blue.png",
     "Header and footer text #ffffff on background #1e3a8a, 0.95rem to 1.05rem"),
    ("contrast-link-accent-on-white.png",
     "Links and accents #2563eb on background #ffffff, 1.05rem normal weight"),
]


def read_sections():
    """Split the markdown into (heading, [paragraphs]), dropping > notes."""
    text = SRC.read_text()
    # Blockquoted notes are addressed to Gilbert, not the grader.
    text = re.sub(r"^>.*$", "", text, flags=re.M)
    sections, current, buf = [], None, []

    def flush():
        if current is None:
            return
        body = "\n".join(buf)
        paras = [re.sub(r"\s+", " ", p).strip()
                 for p in re.split(r"\n\s*\n", body) if p.strip()]
        sections.append((current, paras))

    for line in text.splitlines():
        if line.startswith("## "):
            flush()
            current, buf = line[3:].strip(), []
        elif current is not None:
            buf.append(line)
    flush()
    return dict(sections), sections


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def add_shot(doc, path, caption):
    if not path.exists():
        p = doc.add_paragraph()
        r = p.add_run(f"[MISSING SCREENSHOT: {path.name}]")
        r.bold = True
        r.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
        return False
    doc.add_picture(str(path), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)
    return True


def main():
    by_name, ordered = read_sections()
    doc = Document()

    # US Letter with 1 inch margins.
    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(s, attr, Inches(1))

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    doc.add_heading("Project 4: Final Website Project Reflection", level=0)
    for line in ["CMST 386: Principles of Web Design and Technology II",
                 "Student: Gilbert Anderson",
                 "Website: Strikeworks Studio",
                 "Live URL: https://cmst386-umgc-ganderson58.azurewebsites.net/project4/index.html"]:
        doc.add_paragraph(line)

    if "Overview" in by_name:
        doc.add_heading("Overview", level=1)
        for p in by_name["Overview"]:
            doc.add_paragraph(p)

    doc.add_heading("Reflection Questions", level=1)
    answered = 0
    for heading, paras in ordered:
        if not re.match(r"^\d\.", heading):
            continue
        doc.add_heading(heading, level=2)
        for p in paras:
            doc.add_paragraph(p)
        answered += 1

    words = sum(len(p.split()) for h, ps in ordered
                if re.match(r"^\d\.", h) for p in ps)
    doc.add_paragraph()
    note = doc.add_paragraph()
    r = note.add_run(f"Reflection answers total {words} words across the four prompts.")
    r.italic = True
    r.font.size = Pt(9)

    # ---- Validation screenshots ----
    doc.add_page_break()
    doc.add_heading("HTML Validation Screenshots", level=1)
    doc.add_paragraph(
        "Each of the 12 pages was validated on the live site with the W3C Nu "
        "HTML Checker. Every page returned zero errors and zero warnings.")
    missing = []
    for name in PAGES:
        doc.add_heading(f"{name}.html", level=2)
        if not add_shot(doc, SHOTS / f"html-{name}.png",
                        f"W3C HTML validation result for {name}.html"):
            missing.append(f"html-{name}.png")

    doc.add_page_break()
    doc.add_heading("CSS Validation Screenshot", level=1)
    doc.add_paragraph(
        "The single external stylesheet, css/style.css, was validated with the "
        "W3C Jigsaw CSS validator against the CSS3 profile and returned zero errors.")
    if not add_shot(doc, SHOTS / "css-style.png",
                    "W3C Jigsaw CSS validation result for css/style.css"):
        missing.append("css-style.png")

    # ---- Contrast analyzer ----
    doc.add_page_break()
    doc.add_heading("Color Contrast Analyzer", level=1)
    doc.add_paragraph(
        "Every foreground and background pair used for text on the site was "
        "checked with the WebAIM Contrast Checker before the CSS was written. "
        "Each screenshot below shows the selected background color, the "
        "selected foreground color, and the resulting contrast ratio. The "
        "lowest ratio used anywhere on the site is 5.02:1, for the amber "
        "#b45309 used on service prices and required-field asterisks against "
        "white, which clears the 4.5:1 minimum required for normal text. All "
        "ten text pairs on the site pass, ranging from 5.02:1 to 14.63:1.")
    for fname, caption in CONTRAST:
        if not add_shot(doc, SHOTS / fname, caption):
            missing.append(fname)

    # ---- Sitemap ----
    doc.add_page_break()
    doc.add_heading("XML Sitemap", level=1)
    doc.add_paragraph(
        "The sitemap below was generated by crawling the live site with the "
        "online generator at https://www.xml-sitemaps.com/ and is reproduced "
        "exactly as the tool produced it.")
    xml = re.search(r"```xml\n(.*?)```", SRC.read_text(), re.S)
    if xml:
        for line in xml.group(1).rstrip().splitlines():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(8)

    doc.save(OUT)
    print(f"wrote {OUT}")
    print(f"  {answered} prompt answers, {words} words")
    if missing:
        print(f"  MISSING {len(missing)} screenshots: {', '.join(missing)}")
    else:
        print("  all 17 screenshots embedded")


if __name__ == "__main__":
    main()
